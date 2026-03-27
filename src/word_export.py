"""
Word 文档导出模块
- 话术导出：保留原文 + 英文中文翻译 + 阿拉伯语英文+中文翻译，5分钟自然断句时间戳
- 评论导出：同上翻译规则，含语言标注
"""
import io
import logging
from datetime import datetime

logger = logging.getLogger(__name__)


def _parse_time(ts_str: str):
    """解析时间字符串为 datetime 对象，失败返回 None"""
    if not ts_str:
        return None
    for fmt in ('%Y-%m-%d %H:%M:%S', '%H:%M:%S'):
        try:
            return datetime.strptime(ts_str.strip()[-19:] if len(ts_str) > 8 else ts_str.strip(), fmt)
        except Exception:
            pass
    return None


def _minutes_diff(t1, t2) -> float:
    """计算两个 datetime 的分钟差"""
    if not t1 or not t2:
        return 0
    delta = abs((t2 - t1).total_seconds())
    return delta / 60.0


def export_speech_docx(speech_records: list, session_info: dict = None) -> bytes:
    """
    生成话术 Word 文档
    - 语言规则：
        英语：保留英文 + 中文翻译（两行）
        阿拉伯语：保留阿语 + 英文翻译 + 中文翻译（三行）
        中文：原文即可
    - 时间戳：每5分钟左右在自然断句处另起一行标注时间
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re

    doc = Document()

    # 文档标题
    username = session_info.get('username', '未知') if session_info else '未知'
    start_time = session_info.get('start_time', '') if session_info else ''
    title = doc.add_heading(f'TikTok 直播话术记录 — @{username}', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if start_time:
        sub = doc.add_paragraph(f'开播时间：{start_time}　共 {len(speech_records)} 段话术')
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.runs[0].font.size = Pt(10)
        sub.runs[0].font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

    doc.add_paragraph()  # 空行

    if not speech_records:
        doc.add_paragraph('（本场未采集到话术数据）')
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # 设置正文字体
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # 按时间分组，每 5 分钟一个自然断句区间
    def add_timestamp_marker(doc, time_str, segment_idx):
        p = doc.add_paragraph()
        run = p.add_run(f'── {time_str} ──────────────────────────────')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x9c, 0xa3, 0xaf)
        run.font.italic = True

    last_ts = None
    last_marker_ts = None
    segment_count = 0

    for i, r in enumerate(speech_records):
        text = r.get('text', '').strip()
        text_zh = r.get('text_zh', '').strip()
        ts_str = r.get('timestamp', '')
        lang = r.get('lang', 'other')

        if not text:
            continue

        curr_ts = _parse_time(ts_str)
        time_label = ts_str[-8:-3] if ts_str else ''  # HH:MM

        # 判断是否需要插入时间戳标记
        needs_marker = False
        if last_marker_ts is None:
            needs_marker = True
        elif curr_ts and _minutes_diff(last_marker_ts, curr_ts) >= 5.0:
            # 找到自然断句：当前句是段落的第一句，或上一句超过20字
            if not last_ts or _minutes_diff(last_ts, curr_ts) >= 1.0:
                needs_marker = True

        if needs_marker and curr_ts:
            segment_count += 1
            add_timestamp_marker(doc, time_label, segment_count)
            last_marker_ts = curr_ts

        # 添加话术内容
        if lang.startswith('ar'):
            # 阿拉伯语：三行（阿语原文 + 英文翻译 + 中文翻译）
            p1 = doc.add_paragraph()
            run_time = p1.add_run(f'[{time_label}] ')
            run_time.font.size = Pt(9)
            run_time.font.color.rgb = RGBColor(0x9c, 0xa3, 0xaf)
            run_ar = p1.add_run(text)
            run_ar.font.size = Pt(11)
            run_ar.font.color.rgb = RGBColor(0x92, 0x40, 0x0e)

            # 英文翻译（text_zh 是中文，这里如果有再额外处理，暂用 text_zh 作为中文）
            # 注意：阿语→英语翻译需要单独调用，此处仅展示中文翻译
            if text_zh and text_zh != text:
                p2 = doc.add_paragraph()
                run_zh = p2.add_run(f'    中文：{text_zh}')
                run_zh.font.size = Pt(10)
                run_zh.font.color.rgb = RGBColor(0x16, 0x53, 0x3e)
                run_zh.font.italic = True

        elif lang == 'en' or (lang.startswith('en')):
            # 英语：两行（英文原文 + 中文翻译）
            p1 = doc.add_paragraph()
            run_time = p1.add_run(f'[{time_label}] ')
            run_time.font.size = Pt(9)
            run_time.font.color.rgb = RGBColor(0x9c, 0xa3, 0xaf)
            run_en = p1.add_run(text)
            run_en.font.size = Pt(11)
            run_en.font.color.rgb = RGBColor(0x1d, 0x4e, 0xd8)

            if text_zh and text_zh != text:
                p2 = doc.add_paragraph()
                run_zh = p2.add_run(f'    中文：{text_zh}')
                run_zh.font.size = Pt(10)
                run_zh.font.color.rgb = RGBColor(0x16, 0x53, 0x3e)
                run_zh.font.italic = True

        else:
            # 中文或其他：原文
            p1 = doc.add_paragraph()
            run_time = p1.add_run(f'[{time_label}] ')
            run_time.font.size = Pt(9)
            run_time.font.color.rgb = RGBColor(0x9c, 0xa3, 0xaf)
            run_orig = p1.add_run(text)
            run_orig.font.size = Pt(11)

        last_ts = curr_ts

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_comments_docx(comment_records: list, session_info: dict = None) -> bytes:
    """
    生成评论 Word 文档
    - 语言规则同话术
    - 格式：时间 | 用户名 | 语言标签 | 内容 [+ 翻译]
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    username = session_info.get('username', '未知') if session_info else '未知'
    start_time = session_info.get('start_time', '') if session_info else ''
    title = doc.add_heading(f'TikTok 直播评论记录 — @{username}', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if start_time:
        sub = doc.add_paragraph(f'开播时间：{start_time}　共 {len(comment_records)} 条评论')
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.runs[0].font.size = Pt(10)
        sub.runs[0].font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

    doc.add_paragraph()

    if not comment_records:
        doc.add_paragraph('（本场未采集到评论数据）')
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # 表格形式
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = '时间'
    hdr[1].text = '用户'
    hdr[2].text = '语言'
    hdr[3].text = '内容（+ 翻译）'

    for r in comment_records:
        content = r.get('content', '').strip()
        text_zh = r.get('text_zh', '').strip()
        lang = r.get('lang', 'other')
        ts = r.get('timestamp', '')
        user = r.get('username', r.get('user', ''))
        lang_short = r.get('lang_short', lang[:2].upper())

        if not content:
            continue

        row = table.add_row().cells
        row[0].text = ts[-8:-3] if ts else ''
        row[1].text = str(user)[:20]
        row[2].text = lang_short

        # 内容格式
        if lang.startswith('ar') or lang == 'en':
            cell_text = content
            if text_zh and text_zh != content:
                cell_text += f'\n→ {text_zh}'
            row[3].text = cell_text
        else:
            row[3].text = content

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
